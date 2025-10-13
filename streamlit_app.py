import streamlit as st
import sys
import io
import docx

# 必要なライブラリのインポートを試行
try:
    import google.generativeai as genai
    import docx
except ImportError:
    st.error(
        "必要なライブラリが見つかりません。以下のコマンドでインストールしてください："
    )
    st.code("pip install google-generativeai python-docx")
    st.info(
        "ターミナルでこのコマンドを実行し、アプリを再起動してください。"
    )
    st.stop()

# StreamlitのUI設定
st.title("💬 チャットボットと学びを振り返ろう！")
st.write(
    "記入済みの学習日記フォーマットをDOCS形式でアップロードすると、その内容に関する対話ができます！"
)

# secretsからAPIキーを読み込む
try:
    gemini_api_key = st.secrets["google_api_key"]
    if not gemini_api_key:
        raise KeyError
except KeyError:
    st.error("APIキーがStreamlitのsecretsに設定されていません。")
    st.info(
        "プロジェクトのルートディレクトリに`.streamlit/secrets.toml`ファイルを作成し、"
        "以下の形式でAPIキーを追加してください。\n\n"
        "```toml\n"
        "google_api_key = \"YOUR_API_KEY_HERE\"\n"
        "```"
        "\n`YOUR_API_KEY_HERE`を実際のAPIキーに置き換えてください。"
    )
    st.stop()

# Gemini APIクライアントの初期化
genai.configure(api_key=gemini_api_key)
model = genai.GenerativeModel('gemini-2.5-flash-preview-05-20')

# === ドキュメントアップローダーの追加 ===
uploaded_file = st.file_uploader(
    "ドキュメントをアップロードしてください",
    type=['txt', 'docx']
)

# === メッセージとドキュメント内容を保存するためのセッション状態変数の作成 ===
if "messages" not in st.session_state:
    st.session_state.messages = []
if "document_content" not in st.session_state:
    st.session_state.document_content = None

# アップロードされたファイルを処理する
if uploaded_file is not None:
    # ファイルを一度だけ読み込む
    if st.session_state.document_content is None:
        try:
            if uploaded_file.type == 'text/plain':
                # .txtファイルの場合、UTF-8でデコード
                document_content = uploaded_file.getvalue().decode('utf-8')
            elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # .docxファイルの場合、python-docxで読み込み
                document = docx.Document(uploaded_file)
                paragraphs = [p.text for p in document.paragraphs]
                document_content = "\n".join(paragraphs)
            else:
                st.error("サポートされていないファイル形式です。")
                st.stop()
            
            st.session_state.document_content = document_content
            st.success("ドキュメントが正常にアップロードされました。")
            st.session_state.messages = [] # 新しいドキュメントがアップロードされたらチャット履歴をリセット
            st.info("これで、ドキュメントの内容について質問できます。")
            
            # === ここに最初の質問を生成するロジックを追加 ===
            initial_prompt = f"これからあなたの学習をサポートします。今日の学習日記を拝見しました。\n\nドキュメント:\n{document_content}\n\nまずは、この日の学習で一番印象に残っていることについて教えていただけますか？"
            
            with st.spinner("思考中です..."):
                response = model.generate_content(
                    initial_prompt,
                    stream=False # 最初の一回はストリーミングではなくてよい
                )
            
            assistant_message = response.text
            st.session_state.messages.append({"role": "assistant", "content": assistant_message})

            st.rerun() # UIを再描画してチャットを表示

        except Exception as e:
            st.error(f"ファイルの読み込み中にエラーが発生しました: {e}")

# === チャットUIの表示 ===
if st.session_state.document_content is None:
    st.info("チャットを開始するには、まずドキュメントをアップロードしてください。")
else:
    # 既存のチャットメッセージの表示
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # ユーザー入力のチャットフィールド
    if prompt := st.chat_input("ドキュメントについて質問してください"):
        
        # 最大ラリー数を超えていないかチェック
        if len(st.session_state.messages) >= 10: # ユーザーとアシスタントのペアで5回なので10
            st.info("今日の振り返りはこれで終わりにしましょう。また明日も頑張ってくださいね！")
            st.stop()
            
        # ユーザーのプロンプトを保存して表示
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        try:
            # Gemini APIに渡すためにメッセージ形式を変換
            history = []
            
            # === ここでドキュメントの内容をシステム指示として最初のコンテキストに追加 ===
            # 対話終了と総括・フィードバックの指示を追加
            document_context = f"""
    # Please paste the text defining the chatbot's role here.
    # ここにチャットボットの役割を定義するテキストを貼り付けてください。

 あなたはユーザーが学習している分野の優秀な指導教員であり、孤独の中独学をする成人学習者の自己成長を支援するコーチとしての役割を担う親しみやすいチャットボットです。あなたは、ユーザーが中長期の取り組みにわたって学習意欲を維持できるよう対話を通して支援してください。
１.	チャットボットの役割と振る舞い
・学習者が自分の言葉で学びを振り返り、気づきを深められるように導くことを重視してください。
・学んでいる内容そのものについて、補足的情報を提供する必要はありません。
・必要に応じて、質問を投げかけたり、要約して返したりしながら、学習者が自分の考えや感情に気づけるようサポートしてください。
・対話が単調にならないよう、適度に質問の表現や切り口を変えて話しかけてください。
このような役割・条件を踏まえ、
「最初に目標を尋ね、以降は学習日記に対して4つほどの質問から自然な対話を繰り広げてください。ARCS-Vを意識しつつ学習者の内面を深める支援を行うチャットボット」として振る舞ってください。
２. 対話におけるインストラクション
* ユーザーから日記のドキュメントがアップロードされますので、書き込まれている学習目標と書き込みを踏まえて学習目標を達成できるよう、対話を通じて学習者の学習意欲の維持と学習目標達成をサポートしてください。
* ユーザーの学習日記を読み、共感を示しながら、その日の出来事や感情についてさらに深く掘り下げるような質問をしてください。
* 結論やアドバイスを急ぐのではなく、ユーザー自身が気づきを得られるように対話を導いてください。
* 対話は最大5回のラリー（ユーザーの質問とあなたの応答のペア）で終えるようにしてください。学習目標が設定されました。
３. 前提：ARCS-Vモデルの本文脈での活用について
ARCS-Vモデルは、学習意欲を高めるための拡張版動機づけモデルです。本来、学習教材や研修のコンテンツを作成する際に用いるものですが、独学する学習者が自ら学習にこのモデルを適応できれば、学習意欲の維持や向上に作用できると考えています。その適応の支援の形として、ARCS-Vモデルを元にしたコーチングのような対話を用います。
今回の対話では、特に以下の観点に焦点を置いて対話をしてください。
○関連性Relevance＝学習内容に対する親しみや意義を持たせ、自ら学ぶ姿勢を形成し、学習者に「やりがい」をもたせます。
・学習内容と学習者の経験を結びつけるための質問をする
・学んだ内容の中でどの部分が重要だと思ったか質問する
・学習内容と学習者の目的を結びつけるためにどうすれば良いかを質問する
・学習者がもともと持っていた学習に対する興味・関心を向上させる問いかけをする
・学習者にやりがいを実感してもらうための質問をする
例えば、「今日の学習・活動内容は、ご自身の生活や仕事の中でどのように役立ちそうですか？」
○自信Confidence＝学習過程で成功体験を味わってもらい、その成功が自分の能力や努力によるものだと思わせることで「やればできる」という自信につなげる側面です。
・学習者が「やればできそう」という期待感を抱くことを促すような問いかけをする
・成功体験を学習者が認識し、学習者が自分の能力に対する信頼を高められるよう促す問いかけをする 
例えば、「今日の学びの中で早速実践に移せそうなものはありましたか？」
○意志Volition＝目標を達成するために努力し続けることに関連する行動と態度全般に働きかける側面です。
・この学習を始めた目的を再訪するきっかけとなる問いかけを行う
・目標に対する自分の進捗を考え、必要に応じて計画を見直すよう促す対話を行う
・学習を行う中で抱く感情について自信をモニタリングするきっかけを与える対話おを行う
例えば、「学習を継続するためにどのような工夫ができそうですか？」
４．学習日記受け取り後の対話
・学習者から送られてきた学習日記の内容をもとに、上記の視点を自然に織り交ぜながら、質問やコメントを行ってください。
・質問は一度に何問も送付せず、自然な会話のように1問ずつ投げかけてください。
・上記の観点を1つずつ明示的に分けて質問するのではなく、学習者の話に応じて自然な流れで深掘りしながら、動機づけや課題感を引き出してください。
・対話は親しみやすく、かつ相手が疲れないよう、質問の数は4問前後に抑え、適度な応答を心がけてください。
・設問は一度に全て提示しないで、会話の流れで問いかけるような対話ができるようにしてください。イメージは、コーチングのヒアリングのような会話です。
・学習者が前向きに学びを続けられるよう、励ましや共感も交えた対話を行ってください。
５．対話終了後の流れ
・対話が落ち着いたら、最後に総括とフィードバックを生成して対話をまとめてください。
・対話のまとめとして、記録ように当日の対話のダイアログをドキュメントにそのままコピーアンドペーストできるcsv.形式で出力してください。
・翌日以降もドキュメントがアップロードされた際には、それまでの対話の内容も踏まえて学習目標を達成できるよう、同じように対話を行い、学習者の学習意欲の維持と学習目標達成をサポートしてください。
最後の応答では、必ず対話の終了を告げ、**その日の学習の総括と簡単なフィードバック**を加えてください。

    # Do not remove the following line.
    # この行は削除しないでください。
    \nドキュメント:\n{st.session_state.document_content}"""
            
            history.append({'role': 'user', 'parts': [document_context]})
            
            # === 既存のチャット履歴も追加 ===
            for msg in st.session_state.messages:
                if isinstance(msg, dict) and "role" in msg and "content" in msg:
                    role = "user" if msg["role"] == "user" else "model"
                    history.append({'role': role, 'parts': [msg["content"]]})
            
            # historyが空でないことを確認
            if not history:
                st.warning("チャット履歴が空です。APIリクエストは送信されません。")
                st.stop()
                
            # Gemini APIを使用して応答を生成（ストリーミング）
            response_stream = model.generate_content(
                history,
                stream=True
            )

            # 応答をチャットにストリーミング表示し、セッション状態に保存
            full_response = ""
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                for chunk in response_stream:
                    if chunk.parts:
                        text_part = chunk.parts[0].text
                        full_response += text_part
                        message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown(full_response)
            
            st.session_state.messages.append({"role": "assistant", "content": full_response})

        except Exception as e:
            st.error("エラーが発生しました。詳細はコンソールを確認してください。")
            print(f"エラーの詳細: {e}", file=sys.stderr)
            st.session_state.messages.append({"role": "assistant", "content": "申し訳ありません、応答の生成中にエラーが発生しました。"})
    
    # === Wordファイル出力ボタンの追加 (対話が終了した場合のみ表示) ===
    if len(st.session_state.messages) >= 10:
        doc = docx.Document()
        doc.add_heading('今日の振り返り', 0)
        for message in st.session_state.messages:
            if message["role"] == "user":
                doc.add_paragraph(f"ユーザー: {message['content']}")
            else:
                doc.add_paragraph(f"チャットボット: {message['content']}")

        # メモリ上でdocxファイルを生成
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        st.download_button(
            label="振り返りドキュメントをダウンロード",
            data=doc_io,
            file_name="振り返り.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
